from pathlib import Path
from io import BytesIO

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS = ROOT / "artifacts" / "ai-innovation-circle" / "stakeholder-screenshots"
OUTPUT = ROOT / "docs"
OUTPUT.mkdir(exist_ok=True)

FOREST = "#0d4630"
DEEP_FOREST = "#082c20"
LIME = "#b5dc33"
PALE_LIME = "#f1f7e7"
INK = "#17251f"
MUTED = "#607069"
LINE = "#dbe6df"
WHITE = "#ffffff"

CAPTIONS = {
    "hub-registration-open.png": "Open Hub registration form",
    "hub-registration-validation.png": "Inline form validation",
    "hub-registration-success.png": "Registration confirmation",
    "recurring-rsvp.png": "Recurring meeting invitation",
    "recurring-rsvp-attending.png": "Recurring RSVP after attending",
    "login-page.png": "Magic-link sign-in page",
    "attendee-goal-dialog.png": "Attendee workspace: Create Goal dialog",
    "one-off-rsvp.png": "One-Off invitation and RSVP",
    "one-off-rsvp-attending.png": "One-Off RSVP after attending",
    "mobile-hub-registration.png": "Hub registration on mobile",
    "mobile-one-off-rsvp.png": "One-Off RSVP on mobile",
}


def screenshot(name: str) -> Path:
    path = SCREENSHOTS / name
    if not path.exists():
        raise FileNotFoundError(path)
    return path


def img_dimensions(path: Path):
    with Image.open(path) as image:
        return image.size


def pdf_image(name: str, width_mm: float = 80):
    path = screenshot(name)
    width = width_mm * mm
    source_width, source_height = img_dimensions(path)
    height = width * source_height / source_width
    return RLImage(str(path), width=width, height=height)


def pdf_caption(name: str):
    return Paragraph(
        CAPTIONS[name],
        ParagraphStyle(
            "image-caption",
            fontName="Helvetica",
            fontSize=7.5,
            leading=10,
            textColor=colors.HexColor(MUTED),
            alignment=TA_CENTER,
            spaceBefore=3,
        ),
    )


def pdf_image_card(name: str, width_mm: float = 80):
    return [pdf_image(name, width_mm), pdf_caption(name)]


def set_cell_background(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in kwargs[edge]:
                    element.set(qn("w:{}".format(key)), str(kwargs[edge][key]))


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_docx_text(cell, text, size=10, color=INK, bold=False, alignment=None):
    paragraph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    return paragraph


def add_docx_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor.from_string((FOREST if level == 1 else DEEP_FOREST).replace("#", ""))
    return paragraph


def add_docx_body(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(":", 1)
        run = paragraph.add_run(first + ":")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(FOREST.replace("#", ""))
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        if not run.font.color.rgb:
            run.font.color.rgb = RGBColor.from_string(INK.replace("#", ""))
    return paragraph


def add_docx_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(INK.replace("#", ""))
    return paragraph


def add_docx_image(document, name, width=6.25):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(screenshot(name)), width=Inches(width))
    caption = document.add_paragraph(CAPTIONS[name])
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    for run in caption.runs:
        run.italic = True
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))


def add_docx_image_grid(document, names, widths=(3.05, 3.05)):
    table = document.add_table(rows=1, cols=len(names))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, name in enumerate(names):
        cell = table.cell(0, index)
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_width, source_height = img_dimensions(screenshot(name))
        max_width = Inches(widths[index] - 0.12)
        max_height = Inches(4.15)
        ratio = min(max_width / source_width, max_height / source_height)
        paragraph.add_run().add_picture(str(screenshot(name)), width=int(source_width * ratio), height=int(source_height * ratio))
        caption = cell.add_paragraph(CAPTIONS[name])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))
        set_cell_border(cell, top={"val": "single", "sz": 5, "color": LINE.replace("#", "")},
                        bottom={"val": "single", "sz": 5, "color": LINE.replace("#", "")},
                        left={"val": "single", "sz": 5, "color": LINE.replace("#", "")},
                        right={"val": "single", "sz": 5, "color": LINE.replace("#", "")})
    document.add_paragraph()


def build_docx():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK.replace("#", ""))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Kinetics Group")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(FOREST.replace("#", ""))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Innovation Hubs — attendee registration flows")
    run.font.name = "Aptos Display"
    run.font.size = Pt(19)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DEEP_FOREST.replace("#", ""))

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_after = Pt(14)
    run = intro.add_run("A visual guide to what an attendee sees from first invitation through response.")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))

    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.cell(0, 0)
    set_cell_background(cell, PALE_LIME)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    add_docx_text(
        cell,
        "Executive summary  •  A recurring Hub uses a public interest form first, then a no-login meeting RSVP link. A One-Off event goes directly to a private invitation/RSVP link and does not use public registration.",
        size=10,
        color=DEEP_FOREST,
    )
    document.add_paragraph()
    add_docx_image(document, "hub-registration-open.png", width=5.7)
    note = document.add_paragraph("Representative UI captures use isolated test content; the screens and interactions are the actual application.")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))

    add_docx_heading(document, "1. At a glance", 1)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["", "Recurring Hub", "One-Off event"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = Inches(2.2 if i else 1.55)
        set_cell_background(cell, FOREST)
        set_cell_margins(cell)
        add_docx_text(cell, text, size=10, color=WHITE, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Starting point", "Public registration link", "Direct private invitation"),
        ("Login needed to respond", "No", "No"),
        ("Creates an attendee immediately", "No — creates a pending registration", "Yes — attendee is invited directly"),
        ("RSVP page", "Personalised recurring meeting RSVP", "Personalised One-Off RSVP"),
        ("Full workspace", "Available after magic-link sign-in", "Not part of the One-Off flow"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            set_cell_background(cells[i], "#F7FAF7" if len(table.rows) % 2 == 0 else WHITE)
            add_docx_text(cells[i], value, size=9.5, bold=i == 0, color=FOREST if i == 0 else INK)
            set_cell_border(cells[i], bottom={"val": "single", "sz": 4, "color": LINE.replace("#", "")})
    document.add_paragraph()

    add_docx_heading(document, "2. Recurring Hub registration flow", 1)
    add_docx_body(document, "Step 1 — Open the registration link. An attendee opens a Hub-specific link shared by Kinetics. The page is public and does not ask them to create an account.")
    add_docx_image_grid(document, ["hub-registration-open.png", "hub-registration-validation.png"])
    add_docx_body(document, "Step 2 — Submit interest. The attendee enters Full Name, Work Email, and optional Company, then selects Register Interest. Validation appears inline on the form; there is no popup.")
    add_docx_body(document, "Step 3 — See confirmation. The success card confirms Interest Registered. The attendee is not signed in and no login link is sent at this stage.")
    add_docx_image(document, "hub-registration-success.png", width=5.8)
    add_docx_body(document, "What happens next: the submission stays in the Hub’s pending Public Registrations queue. When the Hub is ready, an administrator can promote the person into a meeting and send the meeting invitation.")

    add_docx_heading(document, "3. Hub meeting invitation and attendee workspace", 1)
    add_docx_body(document, "Step 4 — Respond to the meeting invitation. Once invited to a recurring meeting, the attendee receives a personalised RSVP link. The response page is still no-login: the secure token in the link identifies the invitee.")
    add_docx_image_grid(document, ["recurring-rsvp.png", "recurring-rsvp-attending.png"])
    add_docx_body(document, "The attendee chooses I’ll be there or Can’t make it. The page saves the response, updates the status banner, and allows the attendee to change their response later using the same link.")
    add_docx_body(document, "Step 5 — Optional sign-in for the full Hub workspace. The attendee can sign in through the magic-link page to access their recurring Hub workspace, including meetings, goals, suggestions, and invitations.")
    add_docx_image_grid(document, ["login-page.png", "attendee-goal-dialog.png"])
    add_docx_body(document, "Popup/dialog note: the Create Goal dialog is an authenticated workspace interaction. The public registration and RSVP flows use full-page cards and inline feedback rather than modal popups.")

    add_docx_heading(document, "4. One-Off event flow", 1)
    add_docx_body(document, "Step 1 — Open the direct invitation. A One-Off attendee receives a private, personalised link. There is no public Hub registration form and no login requirement.")
    add_docx_image_grid(document, ["one-off-rsvp.png", "one-off-rsvp-attending.png"])
    add_docx_body(document, "Step 2 — Read the invitation and respond. The page shows the event name, date, any invitation message, and the same two RSVP choices. After submitting, the response banner confirms Attending or Not Attending.")
    add_docx_body(document, "One-Off boundary: the attendee does not receive an automatic sign-in link. The token-only RSVP page is the intended experience for the event.")

    add_docx_heading(document, "5. Mobile experience", 1)
    add_docx_body(document, "Both public entry points are designed for links opened from email on a phone. Cards remain centred, fields remain full width, and RSVP actions stack vertically for easy tapping.")
    add_docx_image_grid(document, ["mobile-hub-registration.png", "mobile-one-off-rsvp.png"])

    add_docx_heading(document, "6. Popup and feedback summary", 1)
    for text in [
        "Hub registration: no popup; inline required-field validation and a dedicated Interest Registered confirmation card.",
        "Recurring RSVP: no popup; selected response state, Saving your response… progress text, and a saved-response banner.",
        "One-Off RSVP: no popup; selected response state and a saved-response banner.",
        "Magic-link sign-in: no popup; the page changes to a Check your inbox confirmation state.",
        "Authenticated attendee workspace: modal dialogs are used for actions such as Create Goal.",
        "Invalid or expired links: a dedicated unavailable card explains that the link can no longer be used.",
    ]:
        add_docx_bullet(document, text)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Kinetics Group Innovation Hubs  •  Attendee experience brief")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))

    path = OUTPUT / "Kinetics_Group_Innovation_Hubs_Attendee_Flows.docx"
    document.save(path)
    return path


class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        canvas.Canvas.__init__(self, *args, **kwargs)
        self.pages = []

    def showPage(self):
        self.pages.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self.pages)
        for state in self.pages:
            self.__dict__.update(state)
            self.draw_footer(page_count)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_footer(self, page_count):
        self.saveState()
        self.setStrokeColor(colors.HexColor(LINE))
        self.setLineWidth(0.5)
        self.line(18 * mm, 14 * mm, 192 * mm, 14 * mm)
        self.setFont("Helvetica", 7.5)
        self.setFillColor(colors.HexColor(MUTED))
        self.drawString(18 * mm, 9 * mm, "Kinetics Group Innovation Hubs  •  Attendee experience brief")
        self.drawRightString(192 * mm, 9 * mm, f"{self.getPageNumber()} / {page_count}")
        self.restoreState()


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=27,
            leading=31, alignment=TA_CENTER, textColor=colors.HexColor(FOREST), spaceAfter=3,
        ),
        "subtitle": ParagraphStyle(
            "subtitle", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=16,
            leading=20, alignment=TA_CENTER, textColor=colors.HexColor(DEEP_FOREST), spaceAfter=6,
        ),
        "tagline": ParagraphStyle(
            "tagline", parent=base["Normal"], fontSize=10.5, leading=15,
            alignment=TA_CENTER, textColor=colors.HexColor(MUTED), spaceAfter=10,
        ),
        "h1": ParagraphStyle(
            "h1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=17,
            leading=21, textColor=colors.HexColor(FOREST), spaceBefore=7, spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "h2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12,
            leading=15, textColor=colors.HexColor(DEEP_FOREST), spaceBefore=5, spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "body", parent=base["BodyText"], fontName="Helvetica", fontSize=9.3,
            leading=13.2, textColor=colors.HexColor(INK), spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "small", parent=base["BodyText"], fontName="Helvetica", fontSize=8,
            leading=10.5, textColor=colors.HexColor(MUTED), spaceAfter=4,
        ),
        "table": ParagraphStyle(
            "table", parent=base["BodyText"], fontName="Helvetica", fontSize=8.1,
            leading=10.5, textColor=colors.HexColor(INK),
        ),
        "table_bold": ParagraphStyle(
            "table_bold", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=8.1,
            leading=10.5, textColor=colors.HexColor(FOREST),
        ),
    }


def comparison_table(styles):
    data = [
        [Paragraph("<b>Experience</b>", styles["table"]), Paragraph("<b>Recurring Hub</b>", styles["table"]), Paragraph("<b>One-Off event</b>", styles["table"])],
        [Paragraph("Starting point", styles["table_bold"]), Paragraph("Public registration link", styles["table"]), Paragraph("Direct private invitation", styles["table"])],
        [Paragraph("Login to respond", styles["table_bold"]), Paragraph("No", styles["table"]), Paragraph("No", styles["table"])],
        [Paragraph("Attendee record", styles["table_bold"]), Paragraph("Pending registration first; promoted later", styles["table"]), Paragraph("Invited attendee record from the start", styles["table"])],
        [Paragraph("RSVP", styles["table_bold"]), Paragraph("Personalised recurring meeting RSVP", styles["table"]), Paragraph("Personalised One-Off RSVP", styles["table"])],
        [Paragraph("Full workspace", styles["table_bold"]), Paragraph("Available after magic-link sign-in", styles["table"]), Paragraph("Not part of the One-Off flow", styles["table"])],
    ]
    table = Table(data, colWidths=[37 * mm, 68 * mm, 68 * mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(FOREST)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#f7faf7")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f7faf7"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(LINE)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def pdf_image_grid(names, width_mm=82):
    columns = []
    for name in names:
        columns.append([pdf_image(name, width_mm), pdf_caption(name)])
    table = Table([columns], colWidths=[width_mm * mm + 3 * mm] * len(names))
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor(LINE)),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor(LINE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def callout(text, styles, background=PALE_LIME):
    table = Table([[Paragraph(text, styles["body"])]], colWidths=[174 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(background)),
        ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor(LINE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def build_pdf():
    path = OUTPUT / "Kinetics_Group_Innovation_Hubs_Attendee_Flows.pdf"
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=19 * mm,
        title="Kinetics Group Innovation Hubs — Attendee registration flows",
        author="Kinetics Group",
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=frame)])

    story = [
        Spacer(1, 13 * mm),
        Paragraph("KINETICS GROUP", styles["title"]),
        Paragraph("Innovation Hubs — attendee registration flows", styles["subtitle"]),
        Paragraph("A visual guide to what an attendee sees from first invitation through response.", styles["tagline"]),
        callout("<b>Executive summary</b>  •  A recurring Hub uses a public interest form first, then a no-login meeting RSVP link. A One-Off event goes directly to a private invitation/RSVP link and does not use public registration.", styles),
        Spacer(1, 8 * mm),
        pdf_image("hub-registration-open.png", 112),
        pdf_caption("hub-registration-open.png"),
        Spacer(1, 2 * mm),
        Paragraph("Representative UI captures use isolated test content; the screens and interactions are the actual application.", styles["small"]),
        PageBreak(),

        Paragraph("1. At a glance", styles["h1"]),
        Paragraph("The two attendee journeys are deliberately separated so recurring Hub membership and one-time event participation remain clear.", styles["body"]),
        comparison_table(styles),
        Spacer(1, 6 * mm),
        callout("<b>The simplest explanation for an attendee:</b> Hub registration says “I’m interested in joining this ongoing forum.” One-Off RSVP says “I’m responding to this specific event invitation.”", styles),
        Spacer(1, 4 * mm),
        Paragraph("2. Recurring Hub registration flow", styles["h1"]),
        Paragraph("<b>Step 1 — Open the registration link.</b> An attendee opens a Hub-specific link shared by Kinetics. The page is public and does not ask them to create an account.", styles["body"]),
        pdf_image_grid(["hub-registration-open.png", "hub-registration-validation.png"]),
        Spacer(1, 4 * mm),
        Paragraph("<b>Step 2 — Submit interest.</b> The attendee enters Full Name, Work Email, and optional Company, then selects Register Interest. Validation appears inline on the form; there is no popup.", styles["body"]),
        Paragraph("<b>Step 3 — See confirmation.</b> The success card confirms Interest Registered. The attendee is not signed in and no login link is sent at this stage.", styles["body"]),
        pdf_image("hub-registration-success.png", 72),
        pdf_caption("hub-registration-success.png"),
        Paragraph("<b>What happens next:</b> the submission stays in the Hub’s pending Public Registrations queue. When the Hub is ready, an administrator can promote the person into a meeting and send the meeting invitation.", styles["body"]),
        PageBreak(),

        Paragraph("3. Hub meeting invitation and workspace", styles["h1"]),
        Paragraph("<b>Step 4 — Respond to the meeting invitation.</b> Once invited to a recurring meeting, the attendee receives a personalised RSVP link. The response page is still no-login: the secure token in the link identifies the invitee.", styles["body"]),
        pdf_image_grid(["recurring-rsvp.png", "recurring-rsvp-attending.png"]),
        Spacer(1, 4 * mm),
        Paragraph("The attendee chooses <b>I’ll be there</b> or <b>Can’t make it</b>. The page saves the response, updates the status banner, and allows the attendee to change their response later using the same link.", styles["body"]),
        Paragraph("<b>Step 5 — Optional sign-in for the full Hub workspace.</b> The attendee can sign in through the magic-link page to access their recurring Hub workspace, including meetings, goals, suggestions, and invitations.", styles["body"]),
        pdf_image_grid(["login-page.png", "attendee-goal-dialog.png"]),
        Spacer(1, 3 * mm),
        callout("<b>Popup/dialog note:</b> the Create Goal dialog is an authenticated workspace interaction. Public registration and RSVP use full-page cards and inline feedback rather than modal popups.", styles, "#f7faf7"),
        PageBreak(),

        Paragraph("4. One-Off event flow", styles["h1"]),
        Paragraph("<b>Step 1 — Open the direct invitation.</b> A One-Off attendee receives a private, personalised link. There is no public Hub registration form and no login requirement.", styles["body"]),
        pdf_image_grid(["one-off-rsvp.png", "one-off-rsvp-attending.png"]),
        Spacer(1, 4 * mm),
        Paragraph("<b>Step 2 — Read the invitation and respond.</b> The page shows the event name, date, any invitation message, and the same two RSVP choices. After submitting, the response banner confirms Attending or Not Attending.", styles["body"]),
        Paragraph("<b>One-Off boundary:</b> the attendee does not receive an automatic sign-in link. The token-only RSVP page is the intended experience for the event.", styles["body"]),
        Spacer(1, 4 * mm),
        Paragraph("5. Mobile experience", styles["h1"]),
        Paragraph("Both public entry points are designed for links opened from email on a phone. Cards remain centred, fields remain full width, and RSVP actions stack vertically for easy tapping.", styles["body"]),
        pdf_image_grid(["mobile-hub-registration.png", "mobile-one-off-rsvp.png"], width_mm=53),
        PageBreak(),

        Paragraph("6. Popup and feedback summary", styles["h1"]),
        *[
            Paragraph(f"•  {text}", styles["body"])
            for text in [
                "Hub registration: no popup; inline required-field validation and a dedicated Interest Registered confirmation card.",
                "Recurring RSVP: no popup; selected response state, Saving your response… progress text, and a saved-response banner.",
                "One-Off RSVP: no popup; selected response state and a saved-response banner.",
                "Magic-link sign-in: no popup; the page changes to a Check your inbox confirmation state.",
                "Authenticated attendee workspace: modal dialogs are used for actions such as Create Goal.",
                "Invalid or expired links: a dedicated unavailable card explains that the link can no longer be used.",
            ]
        ],
        Spacer(1, 8 * mm),
        callout("<b>Takeaway for stakeholders:</b> the attendee experience is intentionally low-friction. Public registration and both RSVP paths work without login; authentication is reserved for the richer recurring Hub workspace after someone has joined.", styles),
        Spacer(1, 9 * mm),
        Paragraph("Prepared as a visual product brief for remote stakeholder review.", styles["small"]),
    ]
    doc.build(story, canvasmaker=NumberedCanvas)
    return path


if __name__ == "__main__":
    docx_path = build_docx()
    pdf_path = build_pdf()
    print(docx_path)
    print(pdf_path)